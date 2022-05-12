#Required for screenshots, time and date and the kayboard inputs
import pyautogui
import keyboard
import datetime
#Needed to control the word document inputs
from docx import Document
from docx.shared import Inches
#Needed for the image recognition section
import cv2
import pytesseract
pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'
import re
#Needed for the flebox uplaod
import tkinter as tk
from tkinter import filedialog as fd
root = tk.Tk()
root.geometry('400x400')

#Opening the document control method
document = Document()


#Defning the class for the keylogger
class keylogger():
    
    #The class that gets called when initialised
    def __init__(self, interval):
        #We pass SEND_REPORT_EVERY to interval
        self.interval = interval

    #Main class
    def callback(self, event):
        #This is the key selected
        key = event.name

        if key == "alt":
            todayDate = datetime.datetime.now()
            today = todayDate.strftime("%d-%m-%y_%I.%M.%S")

            #Final screenshot name
            screenshotName = f'Screenshots\{today}.jpg'

            #Code that takes the screenshot
            myScreenshot = pyautogui.screenshot()
            myScreenshot.save(screenshotName)

            #Adds page header
            document.add_heading(today, 1)
            #Adds the image and then a page break to the code documentation
            document.add_picture(screenshotName, width=Inches(6))

            #This uses tesseract to find all the text on the screenshot
            image = cv2.imread(screenshotName)
            imageText = pytesseract.image_to_string(image)

            ##
            ##
            #The following section defines what text is found and will be added to the documentation
            ##
            ##
            win = tk.Toplevel(root)
            win.title("Description")

            document.add_heading("Copy paragraphs for as many times as required", 5)

            #This is for the creation of functions, methods and classes and must go first
            if "def" in imageText:
                classText = "Here a class is created, the name of the class we are discussing is ______. This class is involved in the ______ portion of the program. And Following is the key information within this class."
                document.add_paragraph(classText)

            #This is for the creation of if-else statements
            if "if" or "else" or "elif" in imageText:
                ifText = "Another section of this screenshot shows the use of an if-else statement, this statement compares ___ and ___ values. This is used to _____."
                document.add_paragraph(ifText)    

            #This is for the creation of for loops
            if "for" in imageText:
                forText = "Another section of this screenshot shows the use of a for statement, this statement compares loops through an array in this case the _____ value. This is used to _____."
                document.add_paragraph(forText)

            #This is for the creation of while loops
            if "while" in imageText:
                whileText = "Another section of this screenshot shows the use of a while loop, this statement compares two values ____ and ____, many times until the result is true/correct.  This is used to _____."
                document.add_paragraph(whileText)

            #This is for the creation of foreach loops
            if "foreach(" in imageText:
                foreachText = "Another section of this screenshot shows the use of a foreach statement, this statement compares loops through an array in this case the _____ value. This is used to _____."
                document.add_paragraph(foreachText)

            #Adding a page break and then saving the document
            document.add_page_break()
            document.save('code_explination.docx')
            
    #Defining the things that start on the call up of the class
    def start(self):
        keyboard.on_release(callback=self.callback)
        keyboard.wait()

class imageReferencing():
    """
    def select_file():
        filetypes = (
            ('files', '*.jpg')
        )
        
        filename = fd.askopenfilename(
            title='Open a file',
            initialdir='/',
            filetypes=filetypes
        )

        showinfo(
            title='Selected File',
            message=filename
        )

    def uploadImage():
        open_button = tk.Button(
            root,
            text = 'Open a File',
            command = imageReferencing.select_file()
        )
        open_button.pack(expand=True)

        root.mainloop()
    """
    
#Starting the program
if __name__ == "__main__":
    userInput = input("Start Code Reference or Refrence an Image? type (code or image)")
    if userInput.lower() == "code":
        keylogger = keylogger(interval=60)
        keylogger.start()
    elif userInput.lower() == "image":
        imageReferencing.uploadImage()

root.mainloop()
